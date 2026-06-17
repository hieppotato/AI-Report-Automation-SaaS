from datetime import datetime, timezone
from io import BytesIO
from typing import Any
from uuid import UUID

from app.core.exceptions import NotFoundError
from app.repositories.report_repo import ReportRepository
from app.schemas.export import ReportExportResponse
from app.services.storage_service import StorageService


class ExportService:
    def __init__(self, report_repo: ReportRepository, storage_service: StorageService) -> None:
        self.report_repo = report_repo
        self.storage_service = storage_service

    def export_report(
        self,
        organization_id: UUID,
        report_id: UUID,
        export_format: str,
        expires_in: int = 3600,
    ) -> ReportExportResponse:
        report = self.report_repo.get_report_by_id(organization_id, report_id)
        if not report:
            raise NotFoundError("Report not found.")

        if export_format == "pdf":
            content = self._build_pdf(report)
            content_type = "application/pdf"
            extension = "pdf"
        else:
            content = self._build_docx(report)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"

        file_name = f"{self._safe_title(report)}-{datetime.now(tz=timezone.utc).strftime('%Y%m%d%H%M%S')}.{extension}"
        stored = self.storage_service.upload_generated_report(organization_id, report_id, file_name, content, content_type)
        signed_url = self.storage_service.create_signed_url(
            stored["file_path"],
            expires_in=expires_in,
            bucket=stored["bucket"],
        )
        self._save_export_metadata(organization_id, report_id, report, export_format, stored["file_path"])
        return ReportExportResponse(
            format=export_format,
            file_path=stored["file_path"],
            signed_url=signed_url,
            expires_in=expires_in,
        )

    def delete_export(
        self,
        organization_id: UUID,
        report_id: UUID,
        export_index: int,
    ) -> None:
        report = self.report_repo.get_report_by_id(organization_id, report_id)
        if not report:
            raise NotFoundError("Report not found.")

        report_json = self._report_json(report)
        exports = report_json.get("exports")
        if not isinstance(exports, list) or export_index < 0 or export_index >= len(exports):
            raise NotFoundError("Export entry not found.")

        removed = exports.pop(export_index)
        file_path = removed.get("file_path")

        # Delete file from storage (best-effort, don't fail if already removed)
        if file_path:
            try:
                self.storage_service.delete_file(file_path, bucket=self.storage_service.generated_bucket)
            except Exception:
                pass  # File may already be deleted

        report_json["exports"] = exports
        self.report_repo.update_report_fields(organization_id, report_id, {"report_json": report_json})

    def _build_pdf(self, report: dict[str, Any]) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, title=report.get("title") or "AI Report")
        styles = getSampleStyleSheet()
        story = [
            Paragraph(report.get("title") or "AI Business Report", styles["Title"]),
            Spacer(1, 18),
            Paragraph(f"Created: {self._created_date(report)}", styles["Normal"]),
            Spacer(1, 36),
            Paragraph("Executive Summary", styles["Heading1"]),
            Paragraph(self._executive_summary(report), styles["BodyText"]),
            PageBreak(),
            Paragraph("KPI Overview", styles["Heading1"]),
            self._kpi_table(report),
            Spacer(1, 18),
            Paragraph("AI Insights", styles["Heading1"]),
        ]
        story.extend(self._paragraph_list(self._insight_lines(report), styles["BodyText"]))
        story.extend([Spacer(1, 18), Paragraph("Anomalies", styles["Heading1"])])
        story.extend(self._paragraph_list(self._anomaly_lines(report), styles["BodyText"]))
        story.extend([Spacer(1, 18), Paragraph("Charts Summary", styles["Heading1"])])
        story.extend(self._paragraph_list(self._chart_lines(report), styles["BodyText"]))

        for item in story:
            if isinstance(item, Table):
                item.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f4f4f5")),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d4d4d8")),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("PADDING", (0, 0), (-1, -1), 8),
                        ]
                    )
                )
        doc.build(story)
        return buffer.getvalue()

    def _build_docx(self, report: dict[str, Any]) -> bytes:
        from docx import Document

        buffer = BytesIO()
        document = Document()
        document.add_heading(report.get("title") or "AI Business Report", 0)
        document.add_paragraph(f"Created: {self._created_date(report)}")
        document.add_heading("Executive Summary", 1)
        document.add_paragraph(self._executive_summary(report))
        document.add_heading("KPI Overview", 1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        for label, value in self._kpis(report):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
        document.add_heading("AI Insights", 1)
        for line in self._insight_lines(report):
            document.add_paragraph(line, style="List Bullet")
        document.add_heading("Anomalies", 1)
        for line in self._anomaly_lines(report):
            document.add_paragraph(line, style="List Bullet")
        document.add_heading("Charts Summary", 1)
        for line in self._chart_lines(report):
            document.add_paragraph(line, style="List Bullet")
        document.save(buffer)
        return buffer.getvalue()

    def _save_export_metadata(self, organization_id: UUID, report_id: UUID, report: dict[str, Any], export_format: str, file_path: str) -> None:
        report_json = self._report_json(report)
        exports = report_json.get("exports")
        if not isinstance(exports, list):
            exports = []
        exports.append(
            {
                "format": export_format,
                "file_path": file_path,
                "created_at": datetime.now(tz=timezone.utc).isoformat(),
            }
        )
        report_json["exports"] = exports
        self.report_repo.update_report_fields(organization_id, report_id, {"report_json": report_json})

    def _kpi_table(self, report: dict[str, Any]):
        from reportlab.platypus import Table

        return Table([["Metric", "Value"], *self._kpis(report)], colWidths=[180, 220])

    def _kpis(self, report: dict[str, Any]) -> list[list[str]]:
        return [
            ["Total Revenue", str(report.get("total_revenue") or "0")],
            ["Total Orders", str(report.get("total_orders") or "0")],
            ["Average Order Value", str(report.get("aov") or "0")],
            ["Repeat Customer Rate", str(report.get("repeat_rate") or "0")],
        ]

    def _executive_summary(self, report: dict[str, Any]) -> str:
        report_json = self._report_json(report)
        return str(report_json.get("executive_summary") or report_json.get("summary") or "No executive summary available.")

    def _insight_lines(self, report: dict[str, Any]) -> list[str]:
        report_json = self._report_json(report)
        lines = []
        for key in ("key_trends", "risks", "opportunities", "recommendations", "insights"):
            lines.extend(self._as_lines(report_json.get(key), label=key.replace("_", " ").title()))
        if not lines:
            lines.extend(self._as_lines(report.get("insights"), label="Insight"))
        return lines or ["No insights available."]

    def _anomaly_lines(self, report: dict[str, Any]) -> list[str]:
        report_json = self._report_json(report)
        return self._as_lines(report_json.get("anomalies") or report.get("anomalies"), label="Anomaly") or ["No anomalies available."]

    def _chart_lines(self, report: dict[str, Any]) -> list[str]:
        report_json = self._report_json(report)
        return self._as_lines(report_json.get("charts") or report.get("charts"), label="Chart") or ["No charts available."]

    def _paragraph_list(self, lines: list[str], style: Any) -> list[Any]:
        from reportlab.platypus import Paragraph, Spacer

        items = []
        for line in lines:
            items.append(Paragraph(f"- {line}", style))
            items.append(Spacer(1, 6))
        return items

    def _as_lines(self, value: Any, label: str) -> list[str]:
        if isinstance(value, list):
            return [self._line_from_item(item, label) for item in value if item]
        if isinstance(value, dict):
            lines = []
            for key, item in value.items():
                if isinstance(item, list):
                    lines.extend(self._as_lines(item, label=str(key).replace("_", " ").title()))
                else:
                    lines.append(self._line_from_item(item, str(key).replace("_", " ").title()))
            return lines
        if isinstance(value, str) and value.strip():
            return [value.strip()]
        return []

    def _line_from_item(self, item: Any, label: str) -> str:
        if isinstance(item, dict):
            title = item.get("title") or item.get("name") or label
            description = item.get("description") or item.get("summary") or item.get("text") or ""
            return f"{title}: {description}".strip(": ")
        return str(item)

    def _report_json(self, report: dict[str, Any]) -> dict[str, Any]:
        value = report.get("report_json") or {}
        return value if isinstance(value, dict) else {}

    def _created_date(self, report: dict[str, Any]) -> str:
        return str(report.get("created_at") or "")

    def _safe_title(self, report: dict[str, Any]) -> str:
        title = str(report.get("title") or "report").strip().lower()
        return "".join(char if char.isalnum() else "-" for char in title).strip("-") or "report"
